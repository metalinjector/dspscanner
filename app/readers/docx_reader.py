"""Безопасное чтение DOCX (Office Open XML / ZIP)."""
from __future__ import annotations

import email
import logging
from email.policy import default
import re
import zipfile
from pathlib import Path
from typing import List, Optional

from app.config import (
    DOCX_MAX_COMPRESSION_RATIO,
    DOCX_MAX_ENTRIES,
    DOCX_MAX_SINGLE_ENTRY_BYTES,
    DOCX_MAX_UNCOMPRESSED_BYTES,
    DOCX_MAX_XML_BYTES,
)
from app.readers.base import BaseReader, ReaderOutcome
from app.readers.doc_formats import read_rtf

# Дочерний логгер "dspscanner": пишет в тот же файл журнала и вкладку "Журнал"
# GUI, что и остальное приложение (см. app/logging_utils.py).
_logger = logging.getLogger("dspscanner.docx_reader")

_WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_TEXT_TAG = f"{_WORD_NS}t"
_BREAK_TAGS = {f"{_WORD_NS}p", f"{_WORD_NS}br", f"{_WORD_NS}tab", f"{_WORD_NS}cr"}
_RELEVANT_XML_PREFIXES = (
    "word/document", "word/header", "word/footer", "word/footnotes",
    "word/endnotes", "word/comments",
)


def _validate_docx_archive(archive: zipfile.ZipFile) -> tuple[bool, Optional[str]]:
    """Проверяет центральный каталог до распаковки данных.

    Защищает python-docx и резервный XML-парсер от ZIP-bomb и аномально
    больших XML-частей. Архив не извлекается на диск.

    Принимает уже открытый ``ZipFile`` — вызывающий код открывает архив
    ровно один раз.
    """
    try:
        infos = archive.infolist()
    except (OSError, zipfile.BadZipFile, zipfile.LargeZipFile) as exc:
        return False, f"Некорректный ZIP-контейнер DOCX: {exc}"

    if len(infos) > DOCX_MAX_ENTRIES:
        return False, f"Слишком много элементов в DOCX: {len(infos)}"

    total_uncompressed = 0
    for info in infos:
        if info.is_dir():
            continue
        total_uncompressed += info.file_size
        if info.file_size > DOCX_MAX_SINGLE_ENTRY_BYTES:
            return False, f"Слишком крупная часть DOCX: {info.filename}"
        compressed = max(1, info.compress_size)
        if info.file_size > 1_000_000 and info.file_size / compressed > DOCX_MAX_COMPRESSION_RATIO:
            return False, f"Подозрительная степень сжатия DOCX: {info.filename}"
        normalized = info.filename.lower().replace("\\", "/")
        if normalized.endswith(".xml") and info.file_size > DOCX_MAX_XML_BYTES:
            return False, f"Слишком крупный XML внутри DOCX: {info.filename}"
    if total_uncompressed > DOCX_MAX_UNCOMPRESSED_BYTES:
        return False, "DOCX слишком велик после распаковки"
    return True, None


class DocxReader(BaseReader):
    name = "docx"

    def extract_text(self, path: Path) -> ReaderOutcome:
        # Открываем ZipFile ровно один раз — раньше архив открывался в
        # is_file_accessible, _validate_docx_archive и _read_manual_xml.
        try:
            archive = zipfile.ZipFile(path)
        except (OSError, zipfile.BadZipFile, zipfile.LargeZipFile) as exc:
            import errno
            if isinstance(exc, OSError) and exc.errno in (errno.EACCES, errno.EPERM):
                return ReaderOutcome(locked=True, error="Файл заблокирован")
            return ReaderOutcome(error=f"Некорректный DOCX-контейнер: {exc}")

        with archive:
            safe, error = _validate_docx_archive(archive)
            if not safe:
                return ReaderOutcome(error=error or "Небезопасный DOCX-контейнер")

            # Прямой XML-парсер линейно обходит содержимое архива и не страдает
            # квадратичным пересозданием table.row.cells в python-docx на больших
            # табличных документах. python-docx остаётся совместимым fallback.
            text = self._read_manual_xml(archive)
            if text:
                # Импортированный через altChunk раздел лежит в отдельной
                # MHTML-части и не виден XML-парсеру. Без склейки содержимое
                # такого раздела терялось в любом документе, где есть хотя бы
                # один обычный абзац. namelist() уже в памяти, поэтому для
                # документов без altChunk проверка ничего не стоит.
                altchunk = self._read_altchunk_parts(archive)
                if altchunk and altchunk not in text:
                    return ReaderOutcome(
                        text=f"{text}\n\n{altchunk}",
                        method_used="manual-xml+altchunk",
                    )
                return ReaderOutcome(text=text, method_used="manual-xml")

        # python-docx и altchunk открывают файл самостоятельно; они остаются
        # fallback-путями после неудачного основного разбора.
        text = self._read_with_python_docx(path)
        if text:
            return ReaderOutcome(text=text, method_used="python-docx-fallback")

        text = self._read_altchunk(path)
        if text:
            return ReaderOutcome(text=text, method_used="altchunk-mht-parser")

        return ReaderOutcome(error="Не удалось извлечь текст из DOCX (файл повреждён?)")

    @staticmethod
    def _read_with_python_docx(path: Path) -> Optional[str]:
        try:
            import docx
        except ImportError:
            return None
        try:
            document = docx.Document(str(path))
            parts: List[str] = []
            parts.extend(para.text for para in document.paragraphs if para.text)
            for table in document.tables:
                for row in table.rows:
                    # python-docx возвращает один и тот же _Cell несколько раз
                    # подряд для горизонтально объединённых ячеек (gridSpan) в
                    # одной строке — без дедупликации текст объединённой ячейки
                    # попадает в вывод многократно и даёт фантомные лишние
                    # совпадения при поиске. Дедупликация намеренно ограничена
                    # одной строкой (не всей таблицей): id() у промежуточных
                    # _Cell-обёрток python-docx может быть переиспользован после
                    # сборки мусора между строками, что на практике приводило
                    # к ложному “дублированию” и потере текста несвязанных ячеек.
                    seen_cells: set[int] = set()
                    for cell in row.cells:
                        if id(cell._tc) in seen_cells:
                            continue
                        seen_cells.add(id(cell._tc))
                        if cell.text:
                            parts.append(cell.text)
            for section in document.sections:
                for header_footer in (section.header, section.footer):
                    try:
                        parts.extend(para.text for para in header_footer.paragraphs if para.text)
                    except Exception:
                        continue
            result = "\n".join(parts).strip()
            return result or None
        except Exception:
            return None

    @staticmethod
    def _read_manual_xml(archive: zipfile.ZipFile) -> Optional[str]:
        try:
            from defusedxml import ElementTree as ET

            names = sorted(
                name for name in archive.namelist()
                if name.lower().replace("\\", "/").endswith(".xml")
                and name.lower().replace("\\", "/").startswith(_RELEVANT_XML_PREFIXES)
            )
            if not names:
                return None

            document_parts: List[str] = []
            for name in names:
                try:
                    info = archive.getinfo(name)
                    if info.file_size > DOCX_MAX_XML_BYTES:
                        _logger.debug(
                            "DOCX: часть %s пропущена (превышен лимит размера XML)",
                            name,
                        )
                        continue
                    root = ET.fromstring(archive.read(name))
                except Exception as exc:  # noqa: BLE001
                    # Одна повреждённая часть (например, comments.xml) не должна
                    # отменять извлечение текста из остальных исправных частей.
                    _logger.debug("DOCX: часть %s пропущена: %s", name, exc)
                    continue
                parts: List[str] = []
                for elem in root.iter():
                    if elem.tag == _TEXT_TAG and elem.text:
                        parts.append(elem.text)
                    elif elem.tag in _BREAK_TAGS:
                        parts.append("\n")
                text = "".join(parts).strip()
                if text:
                    document_parts.append(text)

            result = "\n\n".join(document_parts)
            result = re.sub(r"[ \t]+", " ", result)
            result = re.sub(r"\n{3,}", "\n\n", result).strip()
            return result or None
        except Exception:
            return None

    @staticmethod
    def _read_altchunk(path: Path) -> Optional[str]:
        """Fallback-путь: открывает архив самостоятельно."""
        try:
            with zipfile.ZipFile(path) as archive:
                return DocxReader._read_altchunk_parts(archive)
        except Exception:
            return None

    @staticmethod
    def _read_altchunk_parts(archive: zipfile.ZipFile) -> Optional[str]:
        try:
            names = [
                name for name in archive.namelist()
                if name.lower().startswith("word/")
                and name.lower().endswith((".mht", ".mhtml"))
            ]
            if not names:
                return None
            info = archive.getinfo(names[0])
            if info.file_size > DOCX_MAX_SINGLE_ENTRY_BYTES:
                return None
            data = archive.read(names[0])
        except Exception:
            return None

        try:
            message = email.message_from_bytes(data, policy=default)
            parts = []
            for part in message.walk():
                if part.is_multipart():
                    continue
                payload = part.get_payload(decode=True)
                if not payload:
                    continue
                content_type = part.get_content_type()
                if content_type == "text/html":
                    charset = part.get_content_charset() or "utf-8"
                    html_text = payload.decode(charset, errors="replace")
                    cleaned = re.sub(r"<style[^>]*>[\s\S]*?</style>", " ", html_text, flags=re.IGNORECASE)
                    cleaned = re.sub(r"<script[^>]*>[\s\S]*?</script>", " ", cleaned, flags=re.IGNORECASE)
                    parts.append(re.sub(r"<[^>]+>", " ", cleaned))
                elif content_type == "text/plain":
                    charset = part.get_content_charset() or "utf-8"
                    parts.append(payload.decode(charset, errors="replace"))
                elif content_type == "application/rtf":
                    rtf_text = read_rtf(payload)
                    if rtf_text:
                        parts.append(rtf_text)
            result = re.sub(r"\s+", " ", " ".join(parts)).strip()
            return result if len(result) >= 5 else None
        except Exception:
            return None
